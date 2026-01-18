import os
import re
from datetime import datetime

import numpy as np
import pandas as pd
from docx import Document

from src.standards import calculate_condition, calculate_rotorbar_severity


def _safe_float(value):
    if value is None:
        return None
    s = str(value).strip()
    if not s:
        return None
    s = s.replace(' ', '')
    if ',' in s and '.' in s:
        if re.fullmatch(r'[+-]?\d{1,3}(\.\d{3})+(,\d+)?', s):
            s = s.replace('.', '').replace(',', '.')
        else:
            s = s.replace(',', '')
    elif ',' in s and '.' not in s:
        if re.fullmatch(r'[+-]?\d+(,\d+)?', s):
            s = s.replace(',', '.')
        else:
            s = s.replace(',', '')
    try:
        return float(s)
    except:
        return None


def _iter_all_text(doc):
    for section in doc.sections:
        try:
            for p in section.header.paragraphs:
                t = p.text.strip()
                if t:
                    yield t
        except:
            pass

    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            yield t

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip().replace('\n', ' ')
                if t:
                    yield t


def _extract_report_date(doc, file_path):
    date_line = None
    all_lines = list(_iter_all_text(doc))
    for line in all_lines[:80]:
        if 'date' in line.lower():
            date_line = line
            break

    candidates = [date_line] if date_line else []
    candidates += all_lines[:80]

    date_re = re.compile(r'(\d{1,2})[/-](\d{1,2})[/-](\d{2,4})(?:\s+(\d{1,2}):(\d{2})(?::(\d{2}))?)?')
    for text in candidates:
        if not text:
            continue
        m = date_re.search(text)
        if not m:
            continue

        p1 = int(m.group(1))
        p2 = int(m.group(2))
        y = int(m.group(3))
        if y < 100:
            y += 2000

        h = int(m.group(4) or 0)
        mi = int(m.group(5) or 0)
        sec = int(m.group(6) or 0)

        if p1 > 12 and p2 <= 12:
            d, mo = p1, p2
        elif p2 > 12 and p1 <= 12:
            mo, d = p1, p2
        else:
            mo, d = p1, p2
        try:
            return datetime(y, mo, d, h, mi, sec).date()
        except:
            continue

    try:
        return datetime.fromtimestamp(os.path.getmtime(file_path)).date()
    except:
        return datetime.now().date()


def _extract_bearing_status(doc):
    text = ' '.join(_iter_all_text(doc))
    m = re.search(
        r'\bbearing\b[^\n\r]{0,80}?(normal|ok|good|alarm|warning|bad|rusak|damage|high|critical)',
        text,
        flags=re.IGNORECASE,
    )
    if not m:
        return None
    v = m.group(1).strip().lower()
    if v in {'ok', 'good'}:
        return 'Normal'
    if v in {'warning'}:
        return 'Alarm'
    if v in {'bad', 'rusak', 'damage', 'high', 'critical'}:
        return 'High'
    if v == 'alarm':
        return 'Alarm'
    if v == 'normal':
        return 'Normal'
    return None


def _table_to_matrix(table):
    return [[cell.text.strip().replace('\n', ' ') for cell in row.cells] for row in table.rows]


def _find_row(matrix, startswith_text):
    needle = startswith_text.strip().lower()
    for row in matrix:
        if not row:
            continue
        if str(row[0]).strip().lower() == needle:
            return row
    return None


def _find_first_row_contains(matrix, text):
    needle = text.lower()
    for row in matrix:
        if any(needle in str(cell).lower() for cell in row):
            return row
    return None


def _parse_currents_from_matrix(matrix):
    out = {}
    for label in ['Current 1', 'Current 2', 'Current 3']:
        row = _find_row(matrix, label)
        if row:
            nums = [_safe_float(x) for x in row[1:]]
            nums = [x for x in nums if x is not None]
            if nums:
                out[label] = nums[0]

    dev = _find_row(matrix, '% dev')
    if dev:
        nums = [_safe_float(x) for x in dev[1:]]
        nums = [x for x in nums if x is not None]
        if nums:
            out['Dev Current'] = nums[0]

    return out


def _parse_voltages_from_matrix(matrix):
    out = {}
    for label in ['Voltage 1', 'Voltage 2', 'Voltage 3']:
        row = _find_row(matrix, label)
        if row:
            nums = [_safe_float(x) for x in row[1:]]
            nums = [x for x in nums if x is not None]
            if nums:
                out[label] = nums[0]

    dev = _find_row(matrix, '% dev')
    if dev:
        nums = [_safe_float(x) for x in dev[1:]]
        nums = [x for x in nums if x is not None]
        if nums:
            out['Dev Voltage'] = nums[0]

    return out


def _parse_rotorbar_from_matrix(matrix):
    out = {}
    measured = _find_row(matrix, 'Measured')
    if not measured:
        return out

    header = matrix[1] if len(matrix) > 1 else None
    if header:
        for i, cell in enumerate(header):
            c = re.sub(r'\s+', ' ', str(cell).strip().lower())

            if c in {'upper sb', 'upper sideband', 'upper side band'}:
                v = _safe_float(measured[i])
                if v is not None:
                    out['Upper Sideband'] = v
                continue
            if c in {'lower sb', 'lower sideband', 'lower side band'}:
                v = _safe_float(measured[i])
                if v is not None:
                    out['Lower Sideband'] = v
                continue
            if c in {'rb hlt index', 'rb hlt. index', 'rb hlt indx', 'rb hlt'}:
                v = _safe_float(measured[i])
                if v is not None:
                    out['Rotorbar Health'] = v
                continue
            if c in {'se, fund', 'se fund', 'se,fund', 'se fund.'}:
                v = _safe_float(measured[i])
                if v is not None:
                    out['Se Fund'] = v
                continue
            if c in {'se, harm', 'se harm', 'se,harm', 'se harm.'}:
                v = _safe_float(measured[i])
                if v is not None:
                    out['Se Harm'] = v
                continue
            if c in {'level %', 'level%', 'level percent', 'level % (se harm/se fund)'}:
                v = _safe_float(measured[i])
                if v is not None:
                    out['Rotorbar Level %'] = v
                continue
            if c in {'severity level', 'severity', 'severity lvl'}:
                v = _safe_float(measured[i])
                if v is not None:
                    out['Rotorbar Severity Level'] = int(round(v))
                continue
            if c in {'rotor condition assessment', 'rotor condition'}:
                t = str(measured[i]).strip()
                if t:
                    out['Rotor Condition Assessment'] = t
                continue
            if c in {'recommended corrective action', 'corrective action', 'recommended action'}:
                t = str(measured[i]).strip()
                if t:
                    out['Rotor Corrective Action'] = t
                continue
    return out


def _canon_ps_text(s):
    s = str(s or '').strip()
    s = s.replace('\u00a0', ' ')
    s = re.sub(r'\s+', ' ', s)
    return s


def _translate_performance_choice(text):
    t = _canon_ps_text(text)
    key = t.lower()
    mapping = {
        'this induction motor is operating normally, no action is required.': 'Motor induksi ini beroperasi normal, tidak perlu tindakan.',
        'this induction motor exhibits suspicious operation, trending of the induction motor is warranted.': 'Motor induksi ini menunjukkan indikasi mencurigakan, perlu pemantauan tren.',
        'this induction motor exhibits abnormal indications, action is warranted, now.': 'Motor induksi ini menunjukkan indikasi abnormal, perlu tindakan segera.',

        'power factor exceeds 0.85.': 'Faktor daya melebihi 0,85.',
        'power factor is below 0.85, see detailed report.': 'Faktor daya di bawah 0,85, lihat laporan detail.',

        'current variation is within normal limits.': 'Variasi arus masih dalam batas normal.',
        'current variation is beyond normal limits, see detailed report.': 'Variasi arus melebihi batas normal, lihat laporan detail.',

        'voltage variation is within normal limits.': 'Variasi tegangan masih dalam batas normal.',
        'voltage variation is beyond normal limits, see detailed report.': 'Variasi tegangan melebihi batas normal, lihat laporan detail.',
        'rms voltage differs from nameplate by more than 5%.': 'Tegangan RMS berbeda dari nameplate lebih dari 5%.',

        'load on the induction motor is consistent with nameplate values.': 'Beban motor induksi sesuai dengan nilai nameplate.',
        'load on the induction motor exceeds nameplate values, see detailed report.': 'Beban motor induksi melebihi nilai nameplate, lihat laporan detail.',
        'load on the induction motor is less than 25%.': 'Beban motor induksi kurang dari 25%.',

        'connections are normal.': 'Koneksi normal.',
        'voltage ground reference is not neutral.': 'Referensi ground tegangan bukan netral.',
        'loose connection.': 'Ada koneksi longgar.',

        'rotor bar health is normal.': 'Kesehatan rotor bar normal.',
        'rotor bar health is questionable, see detailed report.': 'Kesehatan rotor bar meragukan, lihat laporan detail.',
        'load is insufficient to determine rotor bar health, at this time.': 'Beban tidak mencukupi untuk menentukan kesehatan rotor bar saat ini.',

        'stator health is normal.': 'Kesehatan stator normal.',
        'stator electrical health is questionable.': 'Kesehatan listrik stator meragukan.',
        'stator mechanical health is questionable.': 'Kesehatan mekanik stator meragukan.',
        'turn to turn short.': 'Terjadi hubung singkat antar lilitan.',

        'dynamic or static eccentricity indications do not exist.': 'Tidak ada indikasi eksentrisitas dinamis maupun statis.',
        'indications of static eccentricity exist .': 'Ada indikasi eksentrisitas statis.',
        'indications of dynamic eccentricity exist.': 'Ada indikasi eksentrisitas dinamis.',

        'there is no evidence of harmonic distortion.': 'Tidak ada bukti distorsi harmonik.',
        'there is evidence of harmonic distortion, see detailed report.': 'Ada bukti distorsi harmonik, lihat laporan detail.',

        'there are no indications of mechanical problems like misalignment or unbalance.': 'Tidak ada indikasi masalah mekanik seperti misalignment atau unbalance.',
        'there are indications of mechanical problems like misalignment / unbalance. perform vibr. survey to identify and correct the cause.': 'Ada indikasi misalignment / unbalance. Lakukan survei vibrasi untuk identifikasi dan perbaikan.',

        'there is no evidence of bearing problem.': 'Tidak ada bukti masalah bearing.',
        'indications of potential bearing problems, perform vibration survey to verify.': 'Ada indikasi potensi masalah bearing, lakukan survei vibrasi untuk verifikasi.',
    }

    if key in mapping:
        return mapping[key]
    key = re.sub(r'\bthis induction motor\b', 'motor induksi ini', key)
    key = key.replace('see detailed report.', 'lihat laporan detail.')
    return _canon_ps_text(key)


def _parse_performance_summary_from_matrix(matrix):
    out = {}
    if not matrix:
        return out

    txt = ' '.join(' '.join(r) for r in matrix).lower()
    if 'performance summary' not in txt:
        return out

    section = None
    section_map = {
        'bottom line': 'Ringkasan Kinerja - Kesimpulan',
        'power factor commentary': 'Ringkasan Kinerja - Faktor Daya',
        'current commentary': 'Ringkasan Kinerja - Arus',
        'voltage commentary': 'Ringkasan Kinerja - Tegangan',
        'load commentary': 'Ringkasan Kinerja - Beban',
        'phase connection commentary': 'Ringkasan Kinerja - Koneksi Fasa',
        'rotor commentary': 'Ringkasan Kinerja - Rotor',
        'stator commentary': 'Ringkasan Kinerja - Stator',
        'rotor/stator air-gap characteristics': 'Ringkasan Kinerja - Air-gap Rotor/Stator',
        'harmonic distortion commentary': 'Ringkasan Kinerja - Distorsi Harmonik',
        'misalignment indications': 'Ringkasan Kinerja - Misalignment/Unbalance',
        'bearing commentary': 'Ringkasan Kinerja - Bearing',
    }

    for row in matrix:
        cells = [_canon_ps_text(c) for c in row if _canon_ps_text(c) != '']
        if not cells:
            continue

        row_join = ' '.join(cells).strip()
        row_key = row_join.lower().rstrip(':')
        if row_key in section_map:
            section = section_map[row_key]
            continue

        if section is None:
            continue

        mark = None
        text = None
        raw_cells = [_canon_ps_text(c) for c in row]
        for i, c in enumerate(raw_cells[:3]):
            if c.strip().lower() in {'x', '✓', '✔', 'check', 'checked'}:
                mark = c
                rest = ' '.join([x for x in raw_cells[i + 1:] if x.strip()])
                text = rest.strip()
                break

        if not mark:
            if raw_cells and raw_cells[0].strip().lower().startswith('x') and len(raw_cells[0].strip()) > 1:
                text = raw_cells[0].strip()[1:].strip()

        if text:
            out[section] = _translate_performance_choice(text)

    return out


def _parse_power_from_matrix(matrix):
    out = {}
    avg = _find_row(matrix, 'Avg/Total')
    if not avg:
        return out

    header = matrix[1] if len(matrix) > 1 else None
    if not header:
        return out

    col_pf = None
    col_kw = None
    for i, cell in enumerate(header):
        c = str(cell).strip().lower()
        if c == 'power factor':
            col_pf = i
        elif c == 'kw':
            col_kw = i

    if col_pf is not None:
        pf = _safe_float(avg[col_pf])
        if pf is not None:
            out['power factor'] = pf
    if col_kw is not None:
        kw = _safe_float(avg[col_kw])
        if kw is not None:
            out['Real Power'] = kw

    return out


def _parse_thd_from_matrix(matrix):
    out = {}
    thd_header = _find_row(matrix, '')
    if not matrix or len(matrix) < 2:
        return out

    header = matrix[1]
    col_all = None
    for i, cell in enumerate(header):
        if str(cell).strip().lower() == 'thd all %':
            col_all = i
            break
    if col_all is None:
        return out

    current_vals = []
    voltage_vals = []

    for row in matrix[2:]:
        if not row:
            continue
        name = str(row[0]).strip().lower()
        val = _safe_float(row[col_all] if col_all < len(row) else None)
        if val is None:
            continue
        if name.startswith('current'):
            current_vals.append(val)
        if name.startswith('voltage'):
            voltage_vals.append(val)

    if current_vals:
        out['THD Current %'] = float(np.nanmax(current_vals))
    if voltage_vals:
        out['THD Voltage %'] = float(np.nanmax(voltage_vals))
    return out


def _parse_rated_from_matrix(matrix):
    out = {}
    for row in matrix:
        if len(row) < 2:
            continue
        k = str(row[0]).strip().lower()
        if k == 'full load current':
            out['Full Load Current'] = _safe_float(row[1])
        if k == 'power factor':
            out['power factor'] = _safe_float(row[1])
        if k == 'voltage':
            out['Voltage Rated'] = _safe_float(row[1])
    return out


def _extract_images(doc, output_dir, file_basename):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)
    
    saved_images = []
    image_count = 0
    
    # Iterate over relationships to find images
    # We use a set to avoid duplicates if multiple rels point to same part
    processed_parts = set()
    
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            if rel.target_part in processed_parts:
                continue
            processed_parts.add(rel.target_part)
            
            ext = os.path.splitext(rel.target_ref)[-1]
            if not ext: 
                ext = ".png"
            
            blob = rel.target_part.blob
            # Filter small images (icons, logos) < 3KB
            if len(blob) < 3 * 1024:
                continue
                
            image_count += 1
            filename = f"{file_basename}_img{image_count}{ext}"
            filepath = os.path.join(output_dir, filename)
            
            # Only write if not exists to save time
            if not os.path.exists(filepath):
                with open(filepath, "wb") as f:
                    f.write(blob)
            
            saved_images.append(filename)
            
    return saved_images


def parse_docx_report(file_path, unit_name='Unknown', voltage_level='Unknown', full_name=None):
    code = os.path.basename(file_path).split('_')[0].split('.')[0].strip().upper()
    doc = Document(file_path)
    dt = _extract_report_date(doc, file_path)
    month = dt.strftime('%b').upper()
    year = int(dt.year)
    
    # Extract Images
    # Target: d:/Project/MCSA/data/images
    # We try to deduce root from file_path, or just hardcode for safety
    # file_path example: d:/Project/MCSA/data/Laporan/...
    # We want d:/Project/MCSA/data/images
    
    # Robust way: use absolute path based on this file location? 
    # This file is src/docx_parser.py. Parent is src. Parent is root.
    # Root + data/images
    root_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    img_dir = os.path.join(root_dir, 'data', 'images')
    
    date_str = dt.strftime('%Y-%m-%d')
    base_name = f"{code}_{date_str}"
    
    try:
        _extract_images(doc, img_dir, base_name)
    except Exception as e:
        print(f"Warning: Failed to extract images for {file_path}: {e}")

    values = {}
    rated = {}

    perf = {}

    bearing_status = _extract_bearing_status(doc)
    if bearing_status:
        values['Bearing'] = bearing_status

    for table in doc.tables:
        matrix = _table_to_matrix(table)
        joined = ' '.join(' '.join(r) for r in matrix).lower()

        if 'summary of rotor bar health' in joined or 'rb hlt index' in joined:
            values.update(_parse_rotorbar_from_matrix(matrix))
            continue
        if 'performance summary' in joined:
            perf.update(_parse_performance_summary_from_matrix(matrix))
            continue
        if _find_row(matrix, 'Current 1') and _find_row(matrix, 'Current 2') and _find_row(matrix, 'Current 3'):
            values.update(_parse_currents_from_matrix(matrix))
            continue
        if _find_row(matrix, 'Voltage 1') and _find_row(matrix, 'Voltage 2') and _find_row(matrix, 'Voltage 3'):
            values.update(_parse_voltages_from_matrix(matrix))
            continue
        if _find_row(matrix, 'Avg/Total') and _find_first_row_contains(matrix, 'Real Power'):
            values.update(_parse_power_from_matrix(matrix))
            continue
        if _find_first_row_contains(matrix, 'Harmonic Distortion Results') and _find_first_row_contains(matrix, 'THD All %'):
            values.update(_parse_thd_from_matrix(matrix))
            continue
        if _find_first_row_contains(matrix, 'Full Load Current'):
            rated.update(_parse_rated_from_matrix(matrix))
            continue

    currents = [values.get('Current 1'), values.get('Current 2'), values.get('Current 3')]
    currents = [c for c in currents if c is not None]
    fla = rated.get('Full Load Current')
    if currents and fla:
        iavg = float(np.mean(currents))
        load_pct = (iavg / float(fla)) * 100.0
        values['Load'] = float(np.clip(load_pct, 0, 100))

    cond_inputs = {
        'Dev Voltage': values.get('Dev Voltage'),
        'Dev Current': values.get('Dev Current'),
        'THD Voltage %': values.get('THD Voltage %'),
        'Upper Sideband': values.get('Upper Sideband'),
        'Lower Sideband': values.get('Lower Sideband'),
        'Rotorbar Health': values.get('Rotorbar Health'),
        'Se Fund': values.get('Se Fund'),
        'Se Harm': values.get('Se Harm'),
        'Rotorbar Level %': values.get('Rotorbar Level %'),
        'Bearing': values.get('Bearing'),
    }
    cond = calculate_condition(cond_inputs)
    values['Kondisi'] = cond.get('Overall', 'Normal')
    if 'Bearing' in values and values['Bearing'] is not None:
        values['Bearing'] = cond.get('Bearing', values['Bearing'])

    rb = calculate_rotorbar_severity(cond_inputs)
    if rb.get('Level') is not None:
        values['Rotorbar Severity Level'] = int(rb['Level'])
    if rb.get('Status') is not None:
        values['Rotorbar'] = str(rb['Status'])

    for k, v in perf.items():
        if v is None:
            continue
        values[k] = v

    rows = []
    for p in [
        'Load',
        'Current 1', 'Current 2', 'Current 3', 'Dev Current',
        'Voltage 1', 'Voltage 2', 'Voltage 3', 'Dev Voltage',
        'power factor',
        'THD Current %', 'THD Voltage %',
        'Real Power',
        'Rotorbar Health',
        'Rotorbar Severity Level',
        'Upper Sideband', 'Lower Sideband',
        'Se Fund', 'Se Harm', 'Rotorbar Level %',
        'Rotorbar',
        'Bearing',
        'Kondisi',
    ]:
        if p in values and values[p] is not None:
            v = values[p]
            rows.append({
                'Equipment': code,
                'Parameter': p,
                'Month': month,
                'Year': year,
                'Month_Name': month,
                'Date': str(dt),
                'Raw_Value': str(v),
                'Value': _safe_float(v),
                'Unit': '',
                'Limit': '',
                'Status': str(v) if p in {'Kondisi', 'Bearing', 'Rotorbar'} else ('Info' if str(p).startswith('Ringkasan Kinerja') else 'Normal'),
                'Unit_Name': unit_name,
                'Voltage_Level': voltage_level,
                'Full_Name': full_name or code
            })

    for p, v in values.items():
        if not str(p).startswith('Ringkasan Kinerja'):
            continue
        if v is None:
            continue
        rows.append({
            'Equipment': code,
            'Parameter': p,
            'Month': month,
            'Year': year,
            'Month_Name': month,
            'Date': str(dt),
            'Raw_Value': str(v),
            'Value': None,
            'Unit': '',
            'Limit': '',
            'Status': 'Info',
            'Unit_Name': unit_name,
            'Voltage_Level': voltage_level,
            'Full_Name': full_name or code,
        })

    return rows


def parse_all_reports(laporan_root, meta_mapping):
    rows = []
    if not os.path.exists(laporan_root):
        return pd.DataFrame()

    for root, _, files in os.walk(laporan_root):
        for fn in files:
            if fn.startswith('~$'):
                continue
            ext = os.path.splitext(fn)[1].lower()
            if ext not in {'.docx', '.docm'}:
                continue
            path = os.path.join(root, fn)
            code = fn.split('_')[0].split('.')[0].strip().upper()
            meta = meta_mapping.get(code, {})
            unit_name = meta.get('Unit', 'Unknown')
            voltage_level = meta.get('Voltage', 'Unknown')
            full_name = meta.get('Full_Name', code)
            try:
                rows.extend(parse_docx_report(path, unit_name=unit_name, voltage_level=voltage_level, full_name=full_name))
            except:
                continue

    if not rows:
        return pd.DataFrame()
    return pd.DataFrame(rows)


def parse_all_reports_with_report(laporan_root, meta_mapping):
    rows = []
    report = {
        'total_files': 0,
        'parsed_files': 0,
        'failed_files': 0,
        'failures': []
    }

    if not os.path.exists(laporan_root):
        return pd.DataFrame(), report

    for root, _, files in os.walk(laporan_root):
        for fn in files:
            if fn.startswith('~$'):
                continue
            ext = os.path.splitext(fn)[1].lower()
            if ext not in {'.docx', '.docm'}:
                continue
            report['total_files'] += 1

            path = os.path.join(root, fn)
            code = fn.split('_')[0].split('.')[0].strip().upper()
            meta = meta_mapping.get(code, {})
            unit_name = meta.get('Unit', 'Unknown')
            voltage_level = meta.get('Voltage', 'Unknown')
            full_name = meta.get('Full_Name', code)
            try:
                rows.extend(parse_docx_report(path, unit_name=unit_name, voltage_level=voltage_level, full_name=full_name))
                report['parsed_files'] += 1
            except Exception as e:
                report['failed_files'] += 1
                report['failures'].append({
                    'file': fn,
                    'error': str(e)[:200]
                })

    if not rows:
        return pd.DataFrame(), report
    return pd.DataFrame(rows), report
